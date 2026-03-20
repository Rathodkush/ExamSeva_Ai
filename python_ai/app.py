from flask import Flask, request, jsonify, render_template, redirect, url_for, flash, session
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_cors import CORS
import json
import io
from PIL import Image, ImageFilter, ImageOps
import pytesseract
import fitz
from ocr_nlp import extract_and_compare, _get_model, SIMILARITY_THRESHOLD
import os
from io import BytesIO
try:
    from docx import Document
except Exception:
    Document = None
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
from datetime import datetime
from pathlib import Path
import re
try:
    from models import db, User, Post, Note
    models_available = True
except Exception as e:
    # Avoid unicode/emoji in logs on Windows consoles with cp1252 encoding
    print(f"[WARN] Models import failed: {e}")
    db = None
    User = None
    Post = None
    Note = None
    models_available = False

app = Flask(__name__, template_folder='templates', static_folder='static')
# Enable CORS for all routes (needed for Node.js backend to call this service)
CORS(app)

# Configuration
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'examseva-secret-key-2024')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get(
    'DATABASE_URL',
    'sqlite:///examseva.db'
)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file

# Initialize extensions
if models_available and db:
    db.init_app(app)
    login_manager = LoginManager()
    login_manager.init_app(app)
    login_manager.login_view = 'login'
    login_manager.login_message = 'Please log in to access this page.'
else:
    login_manager = None

if login_manager:
    @login_manager.user_loader
    def load_user(user_id):
        return User.query.get(int(user_id))
else:
    def load_user(user_id):
        return None

# Thread pool for parallel OCR processing
executor = ThreadPoolExecutor(max_workers=4)

# History persistence for repeated-question detection
DATA_DIR = Path(__file__).resolve().parent / "data"
HISTORY_FILE = DATA_DIR / "history.json"
DATA_DIR.mkdir(parents=True, exist_ok=True)

def load_history():
    if HISTORY_FILE.exists():
        try:
            with open(HISTORY_FILE, "r", encoding="utf-8") as fh:
                return json.load(fh)
        except Exception:
            return []
    return []

def save_history(history):
    try:
        with open(HISTORY_FILE, "w", encoding="utf-8") as fh:
            json.dump(history, fh, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Failed to save history: {e}")

def normalize_for_match(text: str) -> str:
    t = text.lower()
    t = re.sub(r'\s+', ' ', t).strip()
    t = re.sub(r'[^a-z0-9\s]', '', t)
    return t

def clean_extracted_text(text: str) -> str:
    """Remove instruction lines, headers, and non-question content."""
    lines = text.split('\n')
    cleaned_lines = []
    
    # Patterns to skip - only very obvious instruction lines
    skip_patterns = [
        r'^\s*attempt\s+any',
        r'^\s*answer\s+the\s+following',
        r'^\s*answer\s+all',
        r'^\s*section\s*[a-z]?\s*:',
        r'^\s*part\s+[a-z]\s*:',
        r'^\s*total\s+marks',
        r'^\s*time\s+allowed',
        r'^\s*instructions?\s*:',
        r'^[\s\d\.\-:]*$',  # Lines with only numbers, dots, colons, dashes
        r'^\s*\d+\s+marks?\s*$',  # Just mark values
        r'^\s*\(\s*\d+\s*marks?\s*\)\s*$',  # (15 marks) alone
    ]
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 5:
            continue
        
        # Skip lines matching instruction patterns
        skip = False
        for pattern in skip_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                skip = True
                break
        
        if not skip:
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)


def _normalize_ws(s: str) -> str:
    return re.sub(r'\s+', ' ', (s or '')).strip()


def _strip_common_noise_lines(text: str) -> str:
    """Remove common header/footer noise from exam papers (page markers, codes, hashes)."""
    if not text:
        return ""
    lines = [l.strip() for l in text.splitlines()]
    cleaned = []
    for l in lines:
        if not l:
            continue
        low = l.lower()
        # Page markers
        if re.search(r'\bpage\s+\d+\s+of\s+\d+\b', low):
            continue
        if re.search(r'\bturn\s+over\b', low):
            continue
        # Subject/paper boilerplate
        if low.startswith('paper /') or 'subject code' in low or low.startswith('paper code'):
            continue
        if 'total marks' in low or 'hours' in low or 'time' in low and 'allowed' in low:
            continue
        # Long hex-like hashes
        if re.fullmatch(r'[0-9a-f]{16,}', low):
            continue
        # Pure numbers / punctuation
        if re.fullmatch(r'[\d\s\.\-:,()]+', l):
            continue
        cleaned.append(l)
    return "\n".join(cleaned)


def parse_exam_paper(text: str) -> list:
    """
    Parse typical university/board exam paper format:
    - "1. Attempt any three of the following: 15 a. ... b. ..."
    Returns a list of sections with question number and subparts.
    """
    text = _strip_common_noise_lines(text)
    if not text:
        return []

    # Work in a whitespace-normalized string but preserve delimiters
    t = text.replace('\r', '\n')
    t = re.sub(r'[ \t]+', ' ', t)

    # Split on major question numbers like "1," "1." "1)" at line starts
    blocks = re.split(r'(?:^|\n)\s*(\d{1,2})\s*[\).,]\s*', t)
    # re.split keeps separators; pattern -> [pre, qno1, body1, qno2, body2, ...]
    sections = []
    if len(blocks) <= 1:
        blocks = ["", "1", t]

    for i in range(1, len(blocks), 2):
        qno = blocks[i].strip()
        body = blocks[i + 1] if i + 1 < len(blocks) else ""
        body = _normalize_ws(body)
        if len(body) < 20:
            continue

        # Detect instruction like "Attempt any three of the following: 15"
        instruction = None
        marks = None
        m = re.search(r'(attempt\s+any\s+\w+.*?following\s*:?)\s*(\d{1,3})?\s*', body, re.IGNORECASE)
        if m:
            instruction = _normalize_ws(m.group(1))
            if m.group(2):
                marks = int(m.group(2))
        else:
            # Some papers place marks before instruction or on same line like "Attempt any three of the following: 15"
            m2 = re.search(r'(attempt\s+any\s+\w+.*?following\s*:?)\s*[:\-]?\s*(\d{1,3})\b', body, re.IGNORECASE)
            if m2:
                instruction = _normalize_ws(m2.group(1))
                marks = int(m2.group(2))

        # Split subparts a-f inside the body
        # Accept "a." "a)" "a -" etc (OCR may remove newlines/spaces, so be permissive)
        part_splits = re.split(r'(?i)(?:^|\s)(?=([a-f])\s*[\).:\-]\s+)', body)
        parts = []
        if len(part_splits) > 1:
            # part_splits looks like [prefix, 'a', ' a. ...', 'b', ' b. ...', ...]
            prefix = _normalize_ws(part_splits[0])
            for j in range(1, len(part_splits), 2):
                label = (part_splits[j] or '').lower()
                ptxt = part_splits[j + 1] if j + 1 < len(part_splits) else ""
                ptxt = re.sub(r'(?i)^\s*[a-f]\s*[\).:\-]\s*', '', ptxt).strip()
                ptxt = _normalize_ws(ptxt)
                if ptxt:
                    parts.append({"label": label, "text": ptxt})
        else:
            prefix = body

        # Clean prefix by removing instruction duplicate if present
        if instruction:
            prefix = re.sub(re.escape(instruction), '', prefix, flags=re.IGNORECASE).strip()
        prefix = _normalize_ws(prefix)

        sections.append({
            "questionNumber": int(qno) if qno.isdigit() else qno,
            "instruction": instruction,
            "marks": marks,
            "text": prefix if prefix else None,
            "parts": parts
        })

    return sections


def extract_pdf_metadata(text: str) -> dict:
    """Try to extract metadata like subject, university, course and paper details from text.
    This is heuristic-based: look for common labels and fallback to first lines.
    """
    meta = {
        'subject': None,
        'university': None,
        'course': None,
        'paper_details': None,
        'year': None,
        'subject_code': None
    }

    if not text or len(text.strip()) == 0:
        return meta

    # Look for explicit labels
    patterns = {
        'subject': [r'Subject\s*[:\-]\s*(.+)', r'Subject\s*(.+)\s*Paper', r'Subject\s*(.+)'],
        'university': [r'University\s*[:\-]\s*(.+)', r'Institute\s*[:\-]\s*(.+)', r'College\s*[:\-]\s*(.+)'],
        'course': [r'Course\s*[:\-]\s*(.+)', r'Class\s*[:\-]\s*(.+)', r'Program\s*[:\-]\s*(.+)'],
        'paper_details': [r'Paper\s*[:\-]\s*(.+)', r'Exam\s*[:\-]\s*(.+)', r'Marks\s*[:\-]\s*(.+)']
    }

    for key, pats in patterns.items():
        for p in pats:
            m = re.search(p, text, re.IGNORECASE)
            if m:
                meta[key] = m.group(1).strip()
                break

    # If still missing, try to use first 6 lines heuristically
    if not meta['university'] or not meta['subject']:
        lines = [l.strip() for l in text.split('\n') if l.strip()][:6]
        if lines:
            # Assume first line might be university/institute
            if not meta['university']:
                if len(lines[0]) < 120 and len(lines[0]) > 5 and any(w.lower() in lines[0].lower() for w in ['university', 'institute', 'college', 'school', 'department']):
                    meta['university'] = lines[0]
                elif len(lines) >= 2 and any(w.lower() in lines[1].lower() for w in ['university', 'institute', 'college', 'school', 'department']):
                    meta['university'] = lines[1]

            # For subject, look for words like 'Subject', or fallback to lines that look like subject names
            if not meta['subject']:
                for line in lines:
                    if re.search(r'^(Subject|Paper|Course)[:\-]', line, re.IGNORECASE):
                        meta['subject'] = re.sub(r'^(Subject|Paper|Course)[:\-]\s*', '', line, flags=re.I).strip()
                        break
                if not meta['subject'] and len(lines) >= 1:
                    # take second or third line if it looks like subject (few words, title-case)
                    candidate = None
                    if len(lines) >= 2:
                        candidate = lines[1]
                    if not candidate and len(lines) >= 3:
                        candidate = lines[2]
                    if candidate and 3 <= len(candidate.split()) <= 6:
                        meta['subject'] = candidate

    # paper details: try to capture 'Paper X - YEAR - CODE' style lines
    if not meta['paper_details']:
        m = re.search(r'(Paper\s*[:\-]?\s*\w+.*\b\d{4}|Paper\s+\w+|Total\s+Marks\s*[:\-])', text, re.IGNORECASE)
        if m:
            meta['paper_details'] = m.group(0).strip()

    # Extract subject code / paper code lines like:
    # "Paper / Subject Code: 53705 / Linux System Administration"
    # "Code: 53702 / Internet of Things"
    code_line = re.search(r'(?:paper\s*/\s*)?subject\s*code\s*:\s*([0-9A-Za-z]+)\s*/\s*([^\n\r]+)', text, re.IGNORECASE)
    if code_line:
        meta['subject_code'] = code_line.group(1).strip()
        # prefer this as subject if it looks like a subject name
        subj_candidate = code_line.group(2).strip()
        if subj_candidate and (not meta['subject']):
            meta['subject'] = subj_candidate[:200]
    else:
        code_line2 = re.search(r'code\s*:\s*([0-9A-Za-z]+)\s*/\s*([^\n\r]+)', text, re.IGNORECASE)
        if code_line2:
            meta['subject_code'] = code_line2.group(1).strip()
            subj_candidate = code_line2.group(2).strip()
            if subj_candidate and (not meta['subject']):
                meta['subject'] = subj_candidate[:200]

    # Extract year (first reasonable 4-digit year)
    y = re.search(r'\b(20\d{2}|19\d{2})\b', text)
    if y:
        meta['year'] = y.group(1)

    # Final cleanup: ensure short strings
    for k in meta:
        if meta[k] and len(meta[k]) > 200:
            meta[k] = meta[k][:200]

    return meta

# ==================== Authentication Routes ====================

@app.route('/')
def index():
    # Show public home page; authenticated users go to dashboard
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        try:
            role = request.form.get('role', 'student')
            # Only allow 'student' or 'admin' roles via public registration
            if role != 'admin':
                role = 'student'
            username = request.form.get('username', '').strip()
            email = request.form.get('email', '').strip()
            password = request.form.get('password', '')
            confirm_password = request.form.get('confirm_password', '')
            name = request.form.get('name', '').strip()
            phone = request.form.get('phone', '').strip()
            
            # Validation
            if not all([username, email, password, confirm_password, name, phone]):
                flash('All required fields must be filled.', 'danger')
                return render_template('register.html')
            
            if password != confirm_password:
                flash('Passwords do not match.', 'danger')
                return render_template('register.html')
            
            if len(password) < 6:
                flash('Password must be at least 6 characters.', 'danger')
                return render_template('register.html')
            
            if User.query.filter_by(username=username).first():
                flash('Username already exists. Please choose another.', 'danger')
                return render_template('register.html')
            
            if User.query.filter_by(email=email).first():
                flash('Email already registered. Please log in or use another email.', 'danger')
                return render_template('register.html')
            
            # Create user
            user = User(username=username, email=email, name=name, phone=phone, role=role)
            user.set_password(password)
            
            # Role-specific fields (only student-specific handled here)
            if role == 'student':
                user.address = request.form.get('address', '').strip()
                user.standard = request.form.get('standard', '').strip()
                user.university = request.form.get('university', '').strip()
                user.course_type = request.form.get('course_type', '').strip()
            
            db.session.add(user)
            db.session.commit()
            
            flash(f'✓ Account created successfully as {role.upper()}! Please log in.', 'success')
            return redirect(url_for('login'))
        
        except Exception as e:
            db.session.rollback()
            flash(f'Error creating account: {str(e)}', 'danger')
            return render_template('register.html')
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        
        user = User.query.filter_by(username=username).first()
        
        if user and user.check_password(password):
            if not user.is_active:
                flash('Your account has been deactivated. Contact admin.', 'danger')
                return render_template('login.html')
            
            user.last_login = datetime.utcnow()
            db.session.commit()
            
            login_user(user)
            flash(f'✓ Welcome back, {user.name}!', 'success')
            return redirect(url_for('dashboard'))
        
        flash('✗ Invalid username or password.', 'danger')
        return render_template('login.html')
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('✓ You have been logged out.', 'info')
    return redirect(url_for('login'))

# ==================== Dashboard Routes ====================

@app.route('/dashboard')
@login_required
def dashboard():
    if current_user.role == 'admin':
        total_users = User.query.count()
        active_users = User.query.filter_by(is_active=True).count()
        total_posts = Post.query.count()
        total_active_posts = Post.query.filter_by(is_active=True).count()
        total_notes = Note.query.count()
        
        # User breakdown
        students = User.query.filter_by(role='student').count()
        admins = User.query.filter_by(role='admin').count()

        return render_template('admin_dashboard.html',
                             total_users=total_users,
                             active_users=active_users,
                             total_posts=total_posts,
                             total_active_posts=total_active_posts,
                             total_notes=total_notes,
                             students=students,
                             admins=admins)

    else:  # student and other non-admin roles
        available_posts = Post.query.filter_by(is_active=True).count()
        available_notes = Note.query.filter_by(is_active=True).count()

        return render_template('student_dashboard.html',
                             available_posts=available_posts,
                             available_notes=available_notes)

# ==================== Profile Routes ====================

@app.route('/profile')
@login_required
def profile():
    return render_template('profile.html', user=current_user)

@app.route('/profile/edit', methods=['GET', 'POST'])
@login_required
def edit_profile():
    if request.method == 'POST':
        try:
            current_user.name = request.form.get('name', current_user.name).strip()
            current_user.phone = request.form.get('phone', current_user.phone).strip()
            
            # Only student profile fields are editable here
            current_user.address = request.form.get('address', current_user.address).strip()
            current_user.standard = request.form.get('standard', current_user.standard).strip()
            current_user.university = request.form.get('university', current_user.university).strip()
            current_user.course_type = request.form.get('course_type', current_user.course_type).strip()
            
            db.session.commit()
            flash('✓ Profile updated successfully!', 'success')
            return redirect(url_for('profile'))
        
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating profile: {str(e)}', 'danger')
            return render_template('edit_profile.html', user=current_user)
    
    return render_template('edit_profile.html', user=current_user)

# ==================== Admin Routes ====================

@app.route('/admin/users')
@login_required
def admin_users():
    if current_user.role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    
    page = request.args.get('page', 1, type=int)
    users = User.query.paginate(page=page, per_page=20)
    return render_template('admin_users.html', users=users)

@app.route('/admin/users/<int:user_id>/toggle', methods=['POST'])
@login_required
def toggle_user_status(user_id):
    if current_user.role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    
    user = User.query.get(user_id)
    if not user or user.id == current_user.id:
        flash('Cannot modify this user.', 'danger')
        return redirect(url_for('admin_users'))
    
    user.is_active = not user.is_active
    db.session.commit()
    
    status = 'activated' if user.is_active else 'deactivated'
    flash(f'✓ User {user.username} has been {status}!', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/users/<int:user_id>/delete', methods=['POST'])
@login_required
def delete_user(user_id):
    if current_user.role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    
    user = User.query.get(user_id)
    if not user or user.id == current_user.id:
        flash('Cannot delete this user.', 'danger')
        return redirect(url_for('admin_users'))
    
    username = user.username
    db.session.delete(user)
    db.session.commit()
    
    flash(f'✓ User {username} has been deleted!', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/posts')
@login_required
def admin_posts():
    if current_user.role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    
    page = request.args.get('page', 1, type=int)
    posts = Post.query.paginate(page=page, per_page=20)
    return render_template('admin_posts.html', posts=posts)

@app.route('/admin/posts/<int:post_id>/toggle', methods=['POST'])
@login_required
def toggle_post_status(post_id):
    if current_user.role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    
    post = Post.query.get(post_id)
    if not post:
        flash('Post not found.', 'danger')
        return redirect(url_for('admin_posts'))
    
    post.is_active = not post.is_active
    db.session.commit()
    
    status = 'activated' if post.is_active else 'deactivated'
    flash(f'✓ Post has been {status}!', 'success')
    return redirect(url_for('admin_posts'))

@app.route('/admin/posts/<int:post_id>/delete', methods=['POST'])
@login_required
def delete_post(post_id):
    if current_user.role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    
    post = Post.query.get(post_id)
    if not post:
        flash('Post not found.', 'danger')
        return redirect(url_for('admin_posts'))
    
    title = post.title
    db.session.delete(post)
    db.session.commit()
    
    flash(f'✓ Post "{title}" has been deleted!', 'success')
    return redirect(url_for('admin_posts'))

@app.route('/admin/notes')
@login_required
def admin_notes():
    if current_user.role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    
    page = request.args.get('page', 1, type=int)
    notes = Note.query.paginate(page=page, per_page=20)
    return render_template('admin_notes.html', notes=notes)


# ---------------- Report & Search Endpoints ----------------
@app.route('/generate_report', methods=['POST'])
def generate_report():
    """Generate a PDF report with numbered questions, answers and source info, wrapped and paginated."""
    try:
        data = request.get_json() or {}
        groups = data.get('groups', [])
        unique = data.get('unique', [])
        metadata = data.get('metadata', {})

        doc = fitz.open()
        page = doc.new_page()

        # Layout parameters
        left = 72
        right = 72
        top = 72
        bottom = 72
        page_width = page.rect.width
        page_height = page.rect.height
        max_width = page_width - left - right
        line_height = 16

        def new_page():
            p = doc.new_page()
            return p

        def write_paragraph(text, fontsize=10, indent=0):
            """Wrap and write a paragraph; adds a small gap after the paragraph."""
            nonlocal page, y
            if not text:
                y += line_height
                return
            import textwrap
            # conservative chars per line estimate
            chars_per_line = max(int(max_width / (fontsize * 0.6)), 40)
            wrapped = textwrap.wrap(text, width=chars_per_line)
            for i, line in enumerate(wrapped):
                if y + line_height > page_height - bottom:
                    page = new_page()
                    y = top
                page.insert_text((left + indent, y), line, fontsize=fontsize)
                y += line_height
            # extra spacing after paragraph
            y += int(line_height * 0.25)

        # Title and metadata
        title = f"Analysis Report - {str(metadata.get('subject','') or '').strip()}"
        page.insert_text((left, top), title, fontsize=18)
        y = top + 34

        write_paragraph(f"Subject: {metadata.get('subject','N/A')}", fontsize=11)
        if metadata.get('subject_code'):
            write_paragraph(f"Subject Code: {metadata.get('subject_code')}", fontsize=11)
        if metadata.get('university') or metadata.get('institutionName'):
            write_paragraph(f"University / Board: {metadata.get('university') or metadata.get('institutionName')}", fontsize=11)
        if metadata.get('classStandard'):
            write_paragraph(f"Class / Standard: {metadata.get('classStandard')}", fontsize=11)
        if metadata.get('courseType'):
            write_paragraph(f"Course Type: {metadata.get('courseType')}", fontsize=11)
        if metadata.get('semester'):
            write_paragraph(f"Semester: {metadata.get('semester','')}", fontsize=11)
        # Prefer explicit year from metadata, fallback to academicYear
        if metadata.get('year') or metadata.get('academicYear'):
            write_paragraph(f"Year: {metadata.get('year') or metadata.get('academicYear','')}", fontsize=11)

        # Optional sections (allow caller to request only 'repeated' or 'unique')
        sections = metadata.get('sections') if isinstance(metadata.get('sections'), list) else None

        if (sections is None) or ('repeated' in sections):
            # Repeated groups (numbered)
            write_paragraph('Repeated Question Groups:', fontsize=13)
            for idx, g in enumerate(groups, start=1):
                rep = g.get('representative') or (g.get('members') and g.get('members')[0] and g.get('members')[0].get('text')) or f'Group {idx}'
                write_paragraph(f"{idx}. {rep}", fontsize=11, indent=12)
                # small variants
                if g.get('members'):
                    for v_idx, m in enumerate(g.get('members'), start=1):
                        mtext = m.get('text')
                        if mtext and mtext != rep:
                            write_paragraph(f"Variant {v_idx}: {mtext}", fontsize=9, indent=24)
                        # include answers if present on variant
                        if isinstance(m, dict) and m.get('answer'):
                            write_paragraph(f"Answer: {m.get('answer')}", fontsize=9, indent=36)

        if (sections is None) or ('unique' in sections):
            # Unique questions (numbered with answers and source info)
            write_paragraph('Unique Questions:', fontsize=13)
            for i, q in enumerate(unique, start=1):
                txt = q.get('text') if isinstance(q, dict) else str(q)
                write_paragraph(f"{i}. {txt}", fontsize=11, indent=12)
                # If an answer was extracted, include it
                ans = None
                if isinstance(q, dict):
                    ans = q.get('answer')
                if ans:
                    write_paragraph(f"Answer: {ans}", fontsize=10, indent=24)

                # include source info when available (page/snippet)
                if isinstance(q, dict):
                    meta_parts = []
                    if q.get('page'):
                        meta_parts.append(f"page: {q.get('page')}")
                    if q.get('snippet'):
                        meta_parts.append(f"snippet: {q.get('snippet')[:120]}")
                    if meta_parts:
                        write_paragraph(' | '.join(meta_parts), fontsize=9, indent=24)

        pdf_bytes = doc.write()
        # Allow custom title/filename via metadata title
        title_override = metadata.get('title') or metadata.get('subject') or 'analysis_report'
        safe_title = re.sub(r'[^A-Za-z0-9 _-]', '', title_override).strip()[:80]
        filename = f"{safe_title or 'analysis_report'}.pdf"
        return (pdf_bytes, 200, {
            'Content-Type': 'application/pdf',
            'Content-Disposition': f'attachment; filename="{filename}"'
        })
    except Exception as e:
        return jsonify({ 'error': 'Failed to generate report', 'detail': str(e) }), 500


@app.route('/generate_docx', methods=['POST'])
def generate_docx():
    """Generate a .docx file containing repeated groups and unique questions."""
    if Document is None:
        return jsonify({ 'error': 'python-docx not installed on server' }), 500
    try:
        payload = request.get_json() or {}
        groups = payload.get('groups', [])
        unique = payload.get('unique', [])
        title = payload.get('title', 'Analysis Questions')

        doc = Document()
        doc.add_heading(title, level=1)

        doc.add_heading('Repeated Question Groups', level=2)
        if not groups:
            doc.add_paragraph('No repeated groups found.')
        else:
            for idx, g in enumerate(groups):
                rep = g.get('representative') or (g.get('members') and g.get('members')[0] and g.get('members')[0].get('text')) or f'Group {idx+1}'
                doc.add_paragraph(f'{idx+1}. {rep}', style='List Number')
                if g.get('members'):
                    for m in g.get('members'):
                        mtext = m.get("text")
                        doc.add_paragraph(f'- {mtext}', style='List Bullet')
                        # Include answer if variant contains it
                        if isinstance(m, dict) and m.get('answer'):
                            doc.add_paragraph(f'    Answer: {m.get("answer")}', style='Intense Quote')

        doc.add_heading('Unique Questions', level=2)
        if not unique:
            doc.add_paragraph('No unique questions found.')
        else:
            for idx, q in enumerate(unique):
                qtxt = q.get('text') if isinstance(q, dict) else str(q)
                p = doc.add_paragraph(f'{idx+1}. {qtxt}', style='List Number')
                # Include detected answer when available
                if isinstance(q, dict) and q.get('answer'):
                    doc.add_paragraph(f'    Answer: {q.get("answer")}', style='Intense Quote')
                # Include source info when available
                if isinstance(q, dict):
                    meta_parts = []
                    if q.get('page'):
                        meta_parts.append(f"page: {q.get('page')}")
                    if q.get('snippet'):
                        meta_parts.append(f"snippet: {q.get('snippet')[:150]}")
                    if len(meta_parts) > 0:
                        # Add a small paragraph under the question with the meta info
                        doc.add_paragraph('    ' + ' | '.join(meta_parts), style='Intense Quote')

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        # sanitize filename
        safe_title = re.sub(r'[^A-Za-z0-9 _-]', '', title)[:100]
        filename = f"{safe_title or 'analysis_questions'}.docx"
        return (bio.read(), 200, {
            'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'Content-Disposition': f'attachment; filename="{filename}"'
        })
    except Exception as e:
        return jsonify({ 'error': 'Failed to generate docx', 'detail': str(e) }), 500


@app.route('/search_pdf', methods=['POST'])
def search_pdf():
    try:
        payload = request.get_json() or {}
        question = payload.get('question','')
        subject = payload.get('subject')
        file_path = payload.get('filePath')

        if not question:
            return jsonify({ 'found': False, 'message': 'Question required' })

        def search_doc(path):
            try:
                if not path or not os.path.exists(path):
                    return None
                doc = fitz.open(path)

                # First pass: literal substring search (fast)
                for pno in range(len(doc)):
                    try:
                        page = doc.load_page(pno)
                        text = page.get_text()
                        if question.lower() in text.lower()[:2000]:
                            # Find exact locations using PDF text search (returns rects)
                            try:
                                rects = page.search_for(question, quads=False)
                                rect_list = []
                                for r in rects:
                                    rect_list.append({
                                        'x0': r.x0,
                                        'y0': r.y0,
                                        'x1': r.x1,
                                        'y1': r.y1
                                    })
                            except Exception:
                                rect_list = []

                            return { 'found': True, 'filePath': path, 'page': pno+1, 'snippet': text[:200], 'rects': rect_list, 'pageWidth': page.rect.width, 'pageHeight': page.rect.height }
                    except Exception:
                        continue

                # Second pass: semantic (embedding) search if exact match failed
                try:
                    model_bundle = _get_model()
                    model = model_bundle['model']
                    util = model_bundle['util']
                    import numpy as _np
                except Exception as e:
                    # Model not available - skip semantic step
                    doc.close()
                    return None

                q_emb = model.encode([question], convert_to_tensor=True)

                for pno in range(len(doc)):
                    try:
                        page = doc.load_page(pno)
                        text = page.get_text()
                        # split into candidate sentences / chunks
                        cand_chunks = [s.strip() for s in re.split(r'(?<=[.?!])\s+', text) if len(s.strip()) > 30]
                        if not cand_chunks:
                            continue

                        # Limit to avoid huge embedding costs
                        cand_chunks = cand_chunks[:500]

                        cand_embs = model.encode(cand_chunks, batch_size=64, convert_to_tensor=True)
                        sims = util.cos_sim(q_emb, cand_embs).cpu().numpy()[0]
                        best_idx = int(_np.argmax(sims))
                        best_score = float(sims[best_idx])

                        if best_score >= float(SIMILARITY_THRESHOLD):
                            best_chunk = cand_chunks[best_idx]
                            snippet = best_chunk[:200]
                            # Attempt to find rects by searching for the best chunk text
                            try:
                                rects = page.search_for(best_chunk.strip(), quads=False)
                                rect_list = []
                                for r in rects:
                                    rect_list.append({ 'x0': r.x0, 'y0': r.y0, 'x1': r.x1, 'y1': r.y1 })
                            except Exception:
                                rect_list = []

                            doc.close()
                            return { 'found': True, 'filePath': path, 'page': pno+1, 'snippet': snippet, 'rects': rect_list, 'score': best_score, 'pageWidth': page.rect.width, 'pageHeight': page.rect.height }
                    except Exception as e:
                        continue

                doc.close()
            except Exception:
                return None
            return None

        # If a specific file path was requested, search it first
        if file_path:
            res = search_doc(file_path)
            if res:
                return jsonify(res)

        # Otherwise, search all notes available in DB
        notes = Note.query.order_by(Note.created_at.desc()).all()
        for note in notes:
            try:
                if not note.filePath or not os.path.exists(note.filePath):
                    continue
                res = search_doc(note.filePath)
                if res:
                    return jsonify(res)
            except Exception:
                continue

        return jsonify({ 'found': False, 'message': 'No matching page found' })
    except Exception as e:
        return jsonify({ 'error': 'Search failed', 'detail': str(e) }), 500

@app.route('/admin/notes/<int:note_id>/delete', methods=['POST'])
@login_required
def delete_note(note_id):
    if current_user.role != 'admin':
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    
    note = Note.query.get(note_id)
    if not note:
        flash('Note not found.', 'danger')
        return redirect(url_for('admin_notes'))
    
    title = note.title
    db.session.delete(note)
    db.session.commit()
    
    flash(f'✓ Note "{title}" has been deleted!', 'success')
    return redirect(url_for('admin_notes'))

# ==================== OCR & Analysis Routes ====================

@app.route('/health', methods=['GET'])
def health():
    """Lightweight health check (does NOT load ML model to avoid timeouts)."""
    return jsonify({
        "status": "ok",
        "service": "examseva-python-ai",
        "timestamp": datetime.utcnow().isoformat() + "Z"
    }), 200


@app.route('/health/model', methods=['GET'])
def health_model():
    """Heavier health check to verify the embedding model can load and run once."""
    try:
        from ocr_nlp import _get_model
        model_status = "not_checked"
        model_error = None
        try:
            model_bundle = _get_model()
            if model_bundle and model_bundle.get('model'):
                test_text = "Test question"
                test_embedding = model_bundle['model'].encode([test_text], convert_to_tensor=False)
                if test_embedding is not None and len(test_embedding) > 0:
                    model_status = "loaded_and_working"
                else:
                    model_status = "loaded_but_not_working"
            else:
                model_status = "not_loaded"
        except Exception as e:
            model_status = "error"
            model_error = str(e)

        data = {
            "status": "ok",
            "service": "examseva-python-ai",
            "model_status": model_status,
            "timestamp": datetime.utcnow().isoformat() + "Z"
        }
        if model_error:
            data["model_error"] = model_error
        return jsonify(data), 200
    except Exception as e:
        return jsonify({
            "status": "error",
            "service": "examseva-python-ai",
            "error": str(e),
            "timestamp": datetime.utcnow().isoformat() + "Z"
        }), 500

def enhance_image_for_ocr(image: Image.Image) -> Image.Image:
    """Enhance image quality for better OCR."""
    try:
        width, height = image.size
        max_dim = max(width, height)
        if max_dim > 1500:  # Only scale down if too large, not up
            scale = 1500 / max_dim
            new_size = (int(width * scale), int(height * scale))
            image = image.resize(new_size, Image.Resampling.LANCZOS)

        gray = image.convert("L")
        gray = ImageOps.autocontrast(gray)
        gray = gray.filter(ImageFilter.UnsharpMask(radius=1.0, percent=120, threshold=3))

        return gray
    except Exception as e:
        print(f"Error enhancing image for OCR: {e}")
        return image

def run_ocr_with_fallback(img: Image.Image) -> str:
    """Run OCR with fallback configs."""
    try:
        config_primary = "--oem 3 --psm 6"
        txt = pytesseract.image_to_string(img, config=config_primary).strip()

        if len(txt) < 40:
            config_fallback = "--oem 3 --psm 4"
            txt_fallback = pytesseract.image_to_string(img, config=config_fallback).strip()
            if len(txt_fallback) > len(txt):
                txt = txt_fallback
        
        # Try one more fallback if still empty
        if len(txt) < 20:
            config_fallback2 = "--oem 3 --psm 3"
            txt_fallback2 = pytesseract.image_to_string(img, config=config_fallback2).strip()
            if len(txt_fallback2) > len(txt):
                txt = txt_fallback2

        txt = txt.replace("\x0c", "").strip()
        return txt
    except Exception as e:
        print(f"OCR error: {e}")
        return ""

def process_pdf_page(page_num, page):
    """Process a single PDF page with advanced OCR."""
    try:
        # For scanned pages, render to image and run OCR
        pix = page.get_pixmap(dpi=150, matrix=fitz.Matrix(1.0, 1.0))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        enhanced = enhance_image_for_ocr(img)
        txt = run_ocr_with_fallback(enhanced)
        if txt:
            return txt
    except Exception as e:
        print(f"Error processing page {page_num}: {e}")
    return None


def is_garbage_text(text: str) -> bool:
    """Heuristic to detect unreadable / binary-like OCR results (hex dumps, base64, low letter density)."""
    if not text or len(text.strip()) == 0:
        return True
    s = text.strip()
    # long hex sequences (common when OCR reads images of non-text)
    if re.search(r'\b([A-F0-9]{12,})\b', s):
        return True
    # base64-like long strings
    if re.search(r'[A-Za-z0-9+/]{50,}={0,2}', s):
        return True

    letters = sum(1 for c in s if c.isalpha())
    digits = sum(1 for c in s if c.isdigit())
    whitespace = sum(1 for c in s if c.isspace())
    total = max(len(s), 1)
    letter_ratio = letters / total

    # If very few letters, consider garbage
    if letter_ratio < 0.20:
        return True

    return False

def process_image_file(file_bytes, filename):
    """Process an image file with advanced OCR."""
    try:
        image = Image.open(io.BytesIO(file_bytes)).convert('RGB')
        width, height = image.size
        max_dim = max(width, height)
        
        # Resize if too large
        if max_dim > 2000:
            scale = 2000 / max_dim
            new_size = (int(width * scale), int(height * scale))
            image = image.resize(new_size, Image.Resampling.LANCZOS)
        
        # Try different enhancement levels
        txt = ""
        
        # First: standard enhancement
        enhanced = enhance_image_for_ocr(image)
        txt = run_ocr_with_fallback(enhanced)
        
        # If text extraction failed, try aggressive enhancement
        if len(txt) < 30:
            try:
                gray = image.convert("L")
                # Aggressive contrast
                gray = ImageOps.autocontrast(gray, cutoff=10)
                txt = pytesseract.image_to_string(gray, config="--oem 3 --psm 6")
            except:
                pass
        
        # If still no text, try adaptive thresholding
        if len(txt) < 30:
            try:
                gray = image.convert("L")
                gray = gray.filter(ImageFilter.EDGE_ENHANCE_MORE)
                txt = pytesseract.image_to_string(gray, config="--oem 3 --psm 4")
            except:
                pass
        
        if txt:
            return txt.strip()
    except Exception as e:
        print(f"Error processing image {filename}: {e}")
    return None

@app.route('/process', methods=['POST'])
def process():
    """Process uploaded files for question extraction - NO authentication required"""
    start_time = time.time()
    print(f"[DEBUG] /process endpoint called at {time.time()}")
    files = request.files.getlist('files')
    print(f"[DEBUG] Received {len(files)} files")
    
    if not files or len(files) == 0:
        print("[DEBUG] No files uploaded")
        return jsonify({
            "error": "No files uploaded",
            "detail": "Please upload at least one file (PDF or image)",
            "groups": [],
            "unique": [],
            "candidates": []
        }), 400
    
    metadata_raw = request.form.get('metadata', '{}')
    try:
        metadata = json.loads(metadata_raw)
    except:
        metadata = {}

    tesseract_cmd = os.environ.get('TESSERACT_CMD')
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    # Quick metadata-only fast pass if requested by client
    metadata_flag = False
    try:
        metadata_flag = bool(metadata.get('fastMetadata')) if isinstance(metadata, dict) else False
    except Exception:
        metadata_flag = False

    if metadata_flag:
        # Try to extract text quickly from first 1-2 pages of each PDF using PyMuPDF text extraction
        combined_preview = []
        for f in files:
            try:
                filename = f.filename.lower()
                if filename.endswith('.pdf'):
                    pdf_bytes = f.read()
                    try:
                        doc = fitz.open(stream=pdf_bytes, filetype='pdf')
                        max_preview = min(len(doc), 2)
                        preview_texts = []
                        for pno in range(max_preview):
                            try:
                                page = doc.load_page(pno)
                                page_text = page.get_text().strip()
                                if page_text:
                                    preview_texts.append(page_text)
                            except Exception:
                                continue
                        doc.close()
                        if preview_texts:
                            combined_preview.append('\n'.join(preview_texts))
                    except Exception:
                        continue
                else:
                    # For images, fall back to full OCR which is slower; skip in fast path
                    continue
            except Exception:
                continue

        if combined_preview:
            preview_text = '\n'.join(combined_preview)
            meta = extract_pdf_metadata(preview_text)
            return jsonify({ 'groups': [], 'unique': [], 'metadata': meta })
        else:
            # No preview text available - fall through to full processing
            pass

    texts = []
    errors = []
    
    try:
        futures = []
        for f in files:
            try:
                filename = f.filename.lower()
                if filename.endswith('.pdf'):
                    pdf_bytes = f.read()
                    try:
                        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                        max_pages = min(len(doc), 10)
                        for page_num in range(max_pages):
                            try:
                                page = doc[page_num]
                                # Try to get embedded text first (fast and accurate when present)
                                try:
                                    page_text = page.get_text().strip()
                                except Exception:
                                    page_text = ""

                                if page_text and len(page_text) > 80 and not is_garbage_text(page_text):
                                    texts.append(page_text)
                                else:
                                    # Fallback to image OCR for scanned pages
                                    future = executor.submit(process_pdf_page, page_num, page)
                                    futures.append(future)
                            except Exception as pe:
                                errors.append(f"Error on page {page_num}: {str(pe)}")
                        doc.close()
                    except Exception as pdf_e:
                        errors.append(f"Error processing PDF {filename}: {str(pdf_e)}")
                else:
                    file_bytes = f.read()
                    try:
                        future = executor.submit(process_image_file, file_bytes, filename)
                        futures.append(future)
                    except Exception as img_e:
                        errors.append(f"Error processing image {filename}: {str(img_e)}")
            except Exception as file_e:
                errors.append(f"Error with file {f.filename}: {str(file_e)}")
        
        # Collect results from futures
        if futures:
            for future in as_completed(futures):
                try:
                    result = future.result()
                    if result and len(result.strip()) > 0:
                        texts.append(result)
                except Exception as result_e:
                    errors.append(f"Error processing result: {str(result_e)}")

        # Filter out garbage-like OCR outputs (hex dumps, base64, low-letter density)
        filtered_texts = []
        garbage_warnings = []
        for i, t in enumerate(texts):
            if t and not is_garbage_text(t):
                filtered_texts.append(t)
            else:
                garbage_warnings.append({'index': i, 'snippet': (t[:160] if t else '')})

        # Replace texts with filtered_texts for downstream NLP
        texts = filtered_texts

        print(f"OCR completed in {time.time() - start_time:.2f}s, extracted {len(texts)} usable text blocks (warnings={len(garbage_warnings)})")
        if errors:
            print(f"Warnings: {errors}")
        if garbage_warnings:
            print(f"Garbage-like OCR outputs detected on {len(garbage_warnings)} blocks; these were ignored")
        
    except Exception as e:
        print(f"Fatal OCR error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": "OCR processing failed",
            "detail": str(e),
            "groups": [],
            "unique": [],
            "candidates": []
        }), 500

    if not texts or all(len(t.strip()) == 0 for t in texts):
        detail_msg = "OCR could not read the files. Please try: 1) Clearer/higher quality scans (300dpi+), 2) Upload a PDF with embedded text if available, 3) Check text is printed (not handwritten), 4) Try a smaller page range or single file."
        # Add garbage detection info if present
        gw = []
        # Try to include garbage snippets if any
        try:
            gw = [g['snippet'] for g in (locals().get('garbage_warnings') or []) if g.get('snippet')]
        except Exception:
            gw = []

        return jsonify({
            "error": "No text extracted from files",
            "detail": detail_msg,
            "groups": [],
            "unique": [],
            "candidates": [],
            "warnings": errors,
            "garbageSnippets": gw
        }), 400

    if texts:
        combined_raw = " ".join(texts)
        combined = combined_raw
        # Clean extracted text to remove instructions and non-question content
        combined = clean_extracted_text(combined)
        if len(combined) > 50000:
            combined = combined[:50000]
        texts = [combined]

    # Fast mode: return structured questions quickly (target <= ~15s), skip embedding clustering
    fast_mode = False
    try:
        fast_mode = bool(metadata.get('fastMode')) if isinstance(metadata, dict) else False
    except Exception:
        fast_mode = False

    if fast_mode:
        try:
            sections = parse_exam_paper(combined_raw if 'combined_raw' in locals() else combined)
            # Flatten into a list for compatibility with existing UI
            flat_items = []
            for s in sections:
                if s.get("text"):
                    flat_items.append({"id": len(flat_items), "text": s["text"], "keywords": []})
                for p in (s.get("parts") or []):
                    flat_items.append({"id": len(flat_items), "text": f"({p.get('label')}) {p.get('text')}", "keywords": []})

            # Build repeated groups using word/keyword matching (fast, no embeddings)
            from ocr_nlp import get_keywords
            n = len(flat_items)
            parent = list(range(n))
            rank = [0] * n
            def find(x):
                while parent[x] != x:
                    parent[x] = parent[parent[x]]
                    x = parent[x]
                return x
            def union(a, b):
                ra, rb = find(a), find(b)
                if ra == rb:
                    return
                if rank[ra] < rank[rb]:
                    parent[ra] = rb
                elif rank[ra] > rank[rb]:
                    parent[rb] = ra
                else:
                    parent[rb] = ra
                    rank[ra] += 1

            texts = [(_normalize_ws(it.get("text", ""))) for it in flat_items]
            norms = [re.sub(r'\s+', ' ', t.strip().lower()) for t in texts]
            kw_sets = [set(get_keywords(t, max_keywords=10)) for t in texts]

            for i in range(n):
                for j in range(i + 1, n):
                    if not texts[i] or not texts[j]:
                        continue
                    if norms[i] == norms[j]:
                        union(i, j)
                        continue
                    k1, k2 = kw_sets[i], kw_sets[j]
                    if not k1 or not k2:
                        continue
                    inter = len(k1 & k2)
                    uni = len(k1 | k2) or 1
                    jacc = inter / uni
                    # Word-specific matching: group if there is meaningful keyword overlap
                    if inter >= 2 and jacc >= 0.45:
                        union(i, j)

            comp = {}
            for i in range(n):
                r = find(i)
                comp.setdefault(r, []).append(i)
            comps = [c for c in comp.values() if len(c) >= 2]

            groups = []
            grouped_ids = set()
            for gid, c in enumerate(comps):
                rep_i = max(c, key=lambda ii: len(texts[ii]))
                representative = texts[rep_i]
                rep_kw = set(get_keywords(representative, max_keywords=12))
                members = []
                for ii in c:
                    grouped_ids.add(ii)
                    k = kw_sets[ii]
                    inter = len(rep_kw & k)
                    uni = len(rep_kw | k) or 1
                    sim = inter / uni
                    members.append({
                        "id": int(ii),
                        "text": texts[ii],
                        "similarity": float(sim),
                        "keywordOverlap": float(sim),
                        "directKeywordOverlap": float(sim),
                        "answerSimilarity": 0.0,
                        "label": None,
                        "score": float(sim)
                    })
                members.sort(key=lambda m: m["score"], reverse=True)
                groups.append({
                    "groupId": int(gid),
                    "representative": representative,
                    "members": members,
                    "keywords": get_keywords(" ".join([texts[ii] for ii in c]), max_keywords=12),
                    "groupSize": len(members)
                })

            unique_items = []
            for i, it in enumerate(flat_items):
                if i not in grouped_ids:
                    unique_items.append({"id": len(unique_items), "text": texts[i], "keywords": list(kw_sets[i])[:12]})

            response_data = {
                "groups": groups,
                "unique": unique_items,
                "extractedSections": sections,
                "metadata": extract_pdf_metadata(combined_raw if 'combined_raw' in locals() else combined),
                "processingTime": round(time.time() - start_time, 2),
                "fastMode": True
            }
            return jsonify(response_data)
        except Exception as e:
            # fall through to deep analysis if parsing fails
            print(f"[WARN] fastMode parsing failed; falling back to deep analysis: {e}")
    
    nlp_start = time.time()
    try:
        result = extract_and_compare(texts)
        print(f"NLP processing completed in {time.time() - nlp_start:.2f}s")
        print(f"Found {len(result.get('groups', []))} groups and {len(result.get('unique', []))} unique questions")
    except Exception as e:
        print(f"NLP processing error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": "NLP processing failed",
            "detail": str(e),
            "groups": [], 
            "unique": [], 
            "candidates": []
        }), 500
    
    # Ensure result has required structure
    groups = result.get("groups", [])
    unique = result.get("unique", [])
    
    # Format groups to match frontend expectations
    formatted_groups = []
    for g in groups:
        formatted_group = {
            "groupId": g.get("groupId", len(formatted_groups)),
            "representative": g.get("representative", ""),
            "members": g.get("members", []),
            "keywords": g.get("keywords", [])
        }
        formatted_groups.append(formatted_group)
    
    # Format unique questions
    formatted_unique = []
    for u in unique:
        formatted_unique.append({
            "id": u.get("id", len(formatted_unique)),
            "text": u.get("text", ""),
            "keywords": u.get("keywords", [])
        })

    history = load_history()
    hist_map = {item.get("normalized"): item for item in history}

    candidates = result.get("candidates", [])
    if not candidates:
        candidates = []
        for g in formatted_groups:
            candidates.append(g.get("representative"))
            for m in g.get("members", []):
                if isinstance(m, dict):
                    candidates.append(m.get("text", ""))
        for u in formatted_unique:
            candidates.append(u.get("text", ""))

    repeated = []
    now = datetime.utcnow().isoformat() + "Z"
    for c in candidates:
        if not c:
            continue
        norm = normalize_for_match(c)
        if not norm:
            continue
        if norm in hist_map:
            entry = hist_map[norm]
            entry["count"] = entry.get("count", 1) + 1
            entry["lastSeen"] = now
            repeated.append({
                "text": c,
                "count": entry["count"],
                "firstSeen": entry.get("firstSeen"),
                "lastSeen": entry.get("lastSeen")
            })
        else:
            entry = {
                "text": c,
                "normalized": norm,
                "firstSeen": now,
                "lastSeen": now,
                "count": 1
            }
            history.append(entry)
            hist_map[norm] = entry

    save_history(history)

    metadata_extracted = extract_pdf_metadata(combined)

    response_data = {
        "groups": formatted_groups,
        "unique": formatted_unique,
        "candidates": candidates,
        "repeated": repeated,
        "metadata": metadata_extracted,
        "processingTime": round(time.time() - start_time, 2)
    }
    
    print(f"✅ Returning analysis: {len(formatted_groups)} groups, {len(formatted_unique)} unique questions | metadata: {json.dumps(metadata_extracted)[:120]}")
    return jsonify(response_data)


@app.route('/enhance_and_process', methods=['POST'])
def enhance_and_process():
    """Enhance uploaded files (images or PDF pages) and re-run OCR + NLP processing.
    Returns a similar structure to /process but indicates enhancement was applied.
    """
    start_time = time.time()
    files = request.files.getlist('files')
    if not files or len(files) == 0:
        return jsonify({"error": "No files uploaded for enhancement", "detail": "Please upload at least one file (PDF or image)"}), 400

    tesseract_cmd = os.environ.get('TESSERACT_CMD')
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    from image_utils import is_blurry_pil, enhance_pil_image
    from sr_utils import apply_sr_image

    # Parse optional metadata flags (e.g., useSR)
    metadata_raw = request.form.get('metadata', '{}')
    try:
        metadata = json.loads(metadata_raw)
    except Exception:
        metadata = {}

    texts = []
    enhanced_pages = []
    errors = []

    try:
        for f in files:
            filename = f.filename.lower()
            if filename.endswith('.pdf'):
                pdf_bytes = f.read()
                try:
                    doc = fitz.open(stream=pdf_bytes, filetype='pdf')
                    max_pages = min(len(doc), 10)
                    for page_num in range(max_pages):
                        page = doc[page_num]
                        
                        # Optimization: Try direct text extraction first
                        extracted_text = page.get_text().strip()
                        if len(extracted_text) > 100:
                            texts.append(extracted_text)
                            continue
                        
                        # Fallback to OCR if no selectable text
                        pix = page.get_pixmap(dpi=150)
                        img = Image.frombytes('RGB', [pix.width, pix.height], pix.samples)
                        
                        txt = pytesseract.image_to_string(img, config='--psm 6')
                        if txt and len(txt.strip()) > 10:
                            texts.append(txt.strip())
                    doc.close()
                except Exception as e:
                    errors.append(f"Error processing PDF {filename}: {e}")
            else:
                file_bytes = f.read()
                try:
                    img = Image.open(io.BytesIO(file_bytes)).convert('RGB')
                    blurry = False
                    try:
                        blurry = is_blurry_pil(img)
                    except Exception:
                        blurry = False

                    if blurry:
                        enhanced_img = enhance_pil_image(img, do_binarize=True)
                        txt = pytesseract.image_to_string(enhanced_img, config='--psm 6')
                        enhanced_pages.append({'file': filename, 'page': None, 'enhanced': True})
                    else:
                        txt = pytesseract.image_to_string(img, config='--psm 6')
                        if not txt or len(txt.strip()) < 50:
                            enhanced_img = enhance_pil_image(img, do_binarize=True)
                            txt = pytesseract.image_to_string(enhanced_img, config='--psm 6')
                            enhanced_pages.append({'file': filename, 'page': None, 'enhanced': True})

                    if txt and len(txt.strip()) > 10 and not is_garbage_text(txt):
                        texts.append(txt.strip())
                except Exception as ie:
                    errors.append(f"Error processing image {filename}: {ie}")

        # Filter same as /process
        filtered_texts = []
        garbage_warnings = []
        for i, t in enumerate(texts):
            if t and not is_garbage_text(t):
                filtered_texts.append(t)
            else:
                garbage_warnings.append({'index': i, 'snippet': (t[:160] if t else '')})

        texts = filtered_texts

    except Exception as e:
        print(f"Enhancement processing error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": "Enhancement processing failed", "detail": str(e)}), 500

    if not texts:
        detail_msg = "Enhancement attempted but OCR still couldn't extract readable text. Try higher quality scans or manual review."
        gw = [g['snippet'] for g in (locals().get('garbage_warnings') or []) if g.get('snippet')]
        return jsonify({
            "error": "No text extracted after enhancement",
            "detail": detail_msg,
            "groups": [],
            "unique": [],
            "candidates": [],
            "warnings": errors,
            "garbageSnippets": gw,
            "enhancedPages": enhanced_pages
        }), 400

    combined = " ".join(texts)
    combined = clean_extracted_text(combined)
    if len(combined) > 50000:
        combined = combined[:50000]

    try:
        result = extract_and_compare([combined])
    except Exception as e:
        return jsonify({"error": "NLP processing failed after enhancement", "detail": str(e)}), 500

    response = {
        "groups": result.get('groups', []),
        "unique": result.get('unique', []),
        "candidates": result.get('candidates', []),
        "metadata": extract_pdf_metadata(combined),
        "processingTime": round(time.time() - start_time, 2),
        "enhancedPages": enhanced_pages,
        "warnings": errors
    }

    return jsonify(response)

def generate_questions_from_text(text, num_questions=10):
    """Generate quiz questions from text content.

    Improvements:
    - Prefer extracting questions from exam-paper patterns (same questions) when possible.
    - Ensure each generated question starts with a question-tag (What/How/Explain/List/Following...).
    - Keep all content strictly derived from the provided text (no outside facts).
    """
    import re

    def ensure_question_tag(q: str) -> str:
        q = (q or "").strip()
        # Remove leading numbering like "1.", "(a)", "a)", "- "
        q = re.sub(r"^\s*(\(?\d+[\).:-]\s+|\(?[a-zA-Z][\).:-]\s+|-+\s+)", "", q).strip()
        if not q:
            return q
        # Already tagged?
        if re.match(r"^(what|how|explain|list|which|why|define|describe|discuss|compare)\b", q, re.I):
            return q
        low = q.lower()
        if low.startswith("following") or "following" in low:
            return "List " + q
        if "how" in low:
            return "How " + q
        if any(k in low for k in ["explain", "describe", "discuss", "elaborate"]):
            return "Explain " + q
        if any(k in low for k in ["list", "name", "state"]):
            return "List " + q
        if "define" in low:
            return "What is " + q
        return "What " + q

    def sentence_pool(txt: str) -> list:
        sents = re.split(r"[.!?\n]+", txt or "")
        sents = [s.strip() for s in sents if 20 < len(s.strip()) < 320]
        # De-dupe (case-insensitive)
        seen = set()
        out = []
        for s in sents:
            key = re.sub(r"\s+", " ", s).strip().lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    # 1) Try: parse as an exam paper and reuse the same questions
    try:
        sections = parse_exam_paper(text)
    except Exception:
        sections = []

    extracted = []
    if sections:
        for s in sections:
            if s.get("text"):
                extracted.append(s["text"])
            for p in (s.get("parts") or []):
                if p.get("text"):
                    extracted.append(p["text"])

        cleaned = []
        seen = set()
        for raw in extracted:
            q = re.sub(r"\s+", " ", (raw or "")).strip()
            if len(q) < 15:
                continue
            key = q.lower()
            if key in seen:
                continue
            seen.add(key)
            cleaned.append(q)

        if len(cleaned) >= max(4, (int(num_questions) or 10) // 2):
            out = []
            pool = sentence_pool(text)
            for i, base in enumerate(cleaned[: (int(num_questions) or 10)]):
                qt = ensure_question_tag(base)
                if qt and not qt.endswith("?"):
                    qt = qt.rstrip(".:;") + "?"
                # Options strictly from notes: use other snippets as distractors
                distractors = []
                for s in pool:
                    if s.lower() == base.lower():
                        continue
                    distractors.append(re.sub(r"\s+", " ", s).strip()[:90])
                    if len(distractors) >= 3:
                        break
                options = [re.sub(r"\s+", " ", base).strip()[:90]] + distractors
                while len(options) < 4:
                    options.append("Not available from notes")
                out.append({"question": qt[:200], "options": options[:4], "correctAnswer": 0})
            return out[: (int(num_questions) or 10)]

    # 2) Fallback: sentence-based generation with tagging
    sentences = sentence_pool(text)
    key_terms = re.findall(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b", text or "")
    key_terms = list(dict.fromkeys(key_terms))[:20]  # stable unique

    questions = []
    for i, sentence in enumerate(sentences[: (int(num_questions) or 10) * 3]):
        if len(sentence) < 30 or len(sentence) > 250:
            continue

        question_text = ensure_question_tag(sentence.strip())
        if question_text and not question_text.endswith("?"):
            question_text = question_text.rstrip(".:;") + "?"

        options = []
        options.append((sentence[:120]).rstrip(".!?,;:"))
        if key_terms:
            related = key_terms[i % len(key_terms)]
            options.append(f"Related to {related}")
        else:
            options.append("A related concept from the notes")

        # Use another notes sentence as distractor when possible
        if i + 1 < len(sentences):
            options.append((sentences[i + 1][:120]).rstrip(".!?,;:"))
        else:
            options.append("Another point from the notes")

        if key_terms and len(key_terms) > 1:
            different = key_terms[(i + 5) % len(key_terms)]
            options.append(f"Something different - {different}")
        else:
            options.append("A contrasting viewpoint")

        cleaned_options = [opt.rstrip(".!?,;:") for opt in options[:4]]
        questions.append({"question": (question_text or sentence)[:200], "options": cleaned_options, "correctAnswer": 0})
        if len(questions) >= (int(num_questions) or 10):
            break

    # If not enough, create remaining from sentences (tagged)
    while len(questions) < (int(num_questions) or 10) and len(sentences) > len(questions):
        idx = len(questions) % len(sentences)
        sent = sentences[idx].rstrip(".!?,;:")
        qtxt = ensure_question_tag(f"the key concept of: {sent[:100]}")
        if qtxt and not qtxt.endswith("?"):
            qtxt = qtxt.rstrip(".:;") + "?"
        questions.append({
            "question": qtxt[:200],
            "options": ["The main idea", "A supporting detail", "A contrasting point", "An example"],
            "correctAnswer": 0
        })

    return questions[: (int(num_questions) or 10)]

def _extract_text_from_epub(epub_bytes):
    try:
        from ebooklib import epub
        book = epub.read_epub(io.BytesIO(epub_bytes))
        items = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            try:
                payload = item.get_content().decode('utf-8')
                # Strip simple HTML tags
                text = re.sub(r'<[^>]+>', ' ', payload)
                items.append(text)
            except Exception:
                continue
        return '\n'.join(items)
    except Exception:
        return ''


def _read_file_texts(files, page_limit=10, force_ocr=False):
    """Read texts from uploaded files. For PDFs, try to extract embedded text first unless force_ocr=True,
    otherwise fall back to image-based OCR. Returns list of text blocks."""
    texts = []
    tesseract_cmd = os.environ.get('TESSERACT_CMD')
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    for f in files:
        filename = f.filename.lower()
        try:
            if filename.endswith('.pdf'):
                pdf_bytes = f.read()
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                max_pages = min(len(doc), page_limit)
                for page_num in range(max_pages):
                    page = doc[page_num]
                    # Try to extract embedded text first (fast and accurate)
                    txt = ''
                    if not force_ocr:
                        try:
                            txt = page.get_text("text") or ''
                        except Exception:
                            txt = ''
                    # If no embedded text or force_ocr requested, do image OCR
                    if force_ocr or not txt.strip():
                        try:
                            pix = page.get_pixmap(dpi=150)
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            ocr_txt = pytesseract.image_to_string(img, config='--psm 6')
                            if ocr_txt and ocr_txt.strip():
                                texts.append(ocr_txt.strip())
                            elif txt and txt.strip():
                                texts.append(txt.strip())
                        except Exception as e:
                            # Fallback: if embedded text exists use it
                            if txt and txt.strip():
                                texts.append(txt.strip())
                    else:
                        texts.append(txt.strip())
                doc.close()
            elif filename.endswith('.epub'):
                file_bytes = f.read()
                txt = _extract_text_from_epub(file_bytes)
                if txt.strip():
                    texts.append(txt.strip())
            elif filename.endswith('.mobi'):
                # MOBI best-effort: try using ebooklib if it supports, otherwise skip
                try:
                    file_bytes = f.read()
                    txt = _extract_text_from_epub(file_bytes)
                    if txt.strip():
                        texts.append(txt.strip())
                except Exception:
                    continue
            else:
                # Treat as an image
                file_bytes = f.read()
                image = Image.open(io.BytesIO(file_bytes)).convert('RGB')
                txt = pytesseract.image_to_string(image, config='--psm 6')
                if txt.strip():
                    texts.append(txt.strip())
        except Exception as e:
            print(f"Error processing file {filename}: {e}")
            continue
    return texts


def classify_difficulty(text: str) -> str:
    txt = (text or '').lower()
    hard_keywords = ['prove', 'derive', 'calculate', 'show', 'demonstrate', 'verify']
    easy_keywords = ['define', 'state', 'what is', 'list', 'write', 'name']
    if any(k in txt for k in hard_keywords):
        return 'Hard'
    if any(k in txt for k in easy_keywords):
        return 'Easy'
    # length heuristic
    words = len(txt.split())
    if words > 40:
        return 'Hard'
    if words > 15:
        return 'Medium'
    return 'Easy'


def _build_quiz_pdf(questions, title='Quiz', metadata=None):
    doc = fitz.open()
    page = doc.new_page()
    left = 72
    top = 72
    bottom = 72
    page_width = page.rect.width
    page_height = page.rect.height
    max_width = page_width - left - 72
    line_height = 14

    def new_page():
        p = doc.new_page()
        return p

    def write_lines(text, fontsize=10, indent=0):
        nonlocal page, y
        if not text:
            y += line_height
            return
        import textwrap
        chars_per_line = max(int(max_width / (fontsize * 0.6)), 40)
        wrapped = textwrap.wrap(text, width=chars_per_line)
        for line in wrapped:
            if y + line_height > page_height - bottom:
                page = new_page()
                y = top
            page.insert_text((left + indent, y), line, fontsize=fontsize)
            y += line_height
        y += int(line_height * 0.25)

    # Title
    page.insert_text((left, top), title, fontsize=18)
    y = top + 34
    if metadata and metadata.get('subject'):
        write_lines(f"Subject: {metadata.get('subject')}", fontsize=11)
    # Add institution/year/semester/state when provided to the header
    if metadata and metadata.get('institutionName'):
        write_lines(f"Institution: {metadata.get('institutionName')}", fontsize=10)
    if metadata and metadata.get('semester'):
        write_lines(f"Semester: {metadata.get('semester')}", fontsize=10)
    if metadata and metadata.get('year'):
        write_lines(f"Year: {metadata.get('year')}", fontsize=10)
    if metadata and metadata.get('state'):
        write_lines(f"State/Board: {metadata.get('state')}", fontsize=10)

    # Group questions by difficulty
    buckets = {'Easy': [], 'Medium': [], 'Hard': []}
    for q in questions:
        txt = q.get('question') if isinstance(q, dict) else str(q)
        d = classify_difficulty(txt)
        buckets[d].append(q)

    for level in ['Easy', 'Medium', 'Hard']:
        bucket = buckets[level]
        if not bucket:
            continue
        write_lines(f"{level} Questions:", fontsize=13)
        for idx, q in enumerate(bucket, start=1):
            qtxt = q.get('question') if isinstance(q, dict) else str(q)
            write_lines(f"{idx}. {qtxt}", fontsize=11, indent=12)
            if isinstance(q, dict) and q.get('options'):
                for opt_idx, opt in enumerate(q.get('options'), start=1):
                    write_lines(f"   {chr(64+opt_idx)}. {opt}", fontsize=10, indent=24)
            # small space
            y += 4

    pdf_bytes = doc.write()
    safe_title = re.sub(r'[^A-Za-z0-9 _-]', '', (title or 'quiz')).strip()[:80]
    filename = f"{safe_title or 'quiz'}.pdf"
    return (pdf_bytes, filename)


def _build_question_paper_pdf(sections, title='Question Paper', metadata=None, include_answer_key=True, include_marking_scheme=True):
    """
    Build a sectioned question paper PDF.
    sections: [
      { "name": "Section A: MCQs", "instructions": "...", "marksPerQuestion": 1, "questions": [ {question, options?, answerKey?, answerText?} ] }
    ]
    """
    doc = fitz.open()
    page = doc.new_page()
    left = 72
    top = 72
    bottom = 72
    page_width = page.rect.width
    page_height = page.rect.height
    max_width = page_width - left - 72
    line_height = 14

    def new_page():
        p = doc.new_page()
        return p

    def write_lines(text, fontsize=10, indent=0):
        nonlocal page, y
        if text is None:
            y += line_height
            return
        text = str(text)
        if not text.strip():
            y += line_height
            return
        import textwrap
        chars_per_line = max(int(max_width / (fontsize * 0.6)), 40)
        wrapped = textwrap.wrap(text, width=chars_per_line)
        for line in wrapped:
            if y + line_height > page_height - bottom:
                page = new_page()
                y = top
            page.insert_text((left + indent, y), line, fontsize=fontsize)
            y += line_height
        y += int(line_height * 0.25)

    # Title
    page.insert_text((left, top), title, fontsize=18)
    y = top + 34

    # Header metadata
    if metadata:
        if metadata.get('subject'):
            write_lines(f"Subject: {metadata.get('subject')}", fontsize=11)
        if metadata.get('institutionName'):
            write_lines(f"Institution: {metadata.get('institutionName')}", fontsize=10)
        if metadata.get('semester'):
            write_lines(f"Semester: {metadata.get('semester')}", fontsize=10)
        if metadata.get('year'):
            write_lines(f"Year: {metadata.get('year')}", fontsize=10)
        if metadata.get('state'):
            write_lines(f"State/Board: {metadata.get('state')}", fontsize=10)
        if metadata.get('examDuration'):
            write_lines(f"Duration: {metadata.get('examDuration')}", fontsize=10)
        if metadata.get('totalMarks'):
            write_lines(f"Total Marks: {metadata.get('totalMarks')}", fontsize=10)

    y += 8

    # Sections
    q_global = 1
    answer_items = []
    scheme_items = []

    for s in sections or []:
        name = s.get("name") or "Section"
        instr = s.get("instructions")
        mpq = s.get("marksPerQuestion")
        qs = s.get("questions") or []

        write_lines(name, fontsize=13)
        if instr:
            write_lines(instr, fontsize=10)
        if mpq:
            write_lines(f"Marks per question: {mpq}", fontsize=10)
        y += 4

        for idx, q in enumerate(qs, start=1):
            qtext = q.get("question") if isinstance(q, dict) else str(q)
            marks = q.get("marks")
            marks_txt = f" ({marks} marks)" if marks else ""
            write_lines(f"{q_global}. {qtext}{marks_txt}", fontsize=11, indent=6)

            opts = q.get("options") if isinstance(q, dict) else None
            if opts:
                for opt_idx, opt in enumerate(opts, start=1):
                    write_lines(f"   {chr(64+opt_idx)}. {opt}", fontsize=10, indent=18)

            # Collect answer key + marking scheme data for later
            if include_answer_key:
                if q.get("answerKey"):
                    answer_items.append((q_global, q.get("answerKey")))
                elif q.get("correctAnswer") is not None and opts:
                    try:
                        ai = int(q.get("correctAnswer"))
                        answer_items.append((q_global, chr(65 + ai)))
                    except Exception:
                        pass
                elif q.get("answerText"):
                    answer_items.append((q_global, q.get("answerText")))

            if include_marking_scheme:
                if q.get("marking"):
                    scheme_items.append((q_global, q.get("marking")))
                elif q.get("answerText"):
                    scheme_items.append((q_global, "Key points from notes: " + str(q.get("answerText"))[:160]))

            y += 3
            q_global += 1

        y += 8

    # Answer key
    if include_answer_key and answer_items:
        write_lines("Answer Key", fontsize=13)
        for qno, ans in answer_items:
            write_lines(f"{qno}. {ans}", fontsize=10, indent=6)
        y += 8

    # Marking scheme
    if include_marking_scheme and scheme_items:
        write_lines("Marking Scheme", fontsize=13)
        for qno, scheme in scheme_items:
            write_lines(f"{qno}. {scheme}", fontsize=10, indent=6)
        y += 8

    pdf_bytes = doc.write()
    safe_title = re.sub(r'[^A-Za-z0-9 _-]', '', (title or 'question_paper')).strip()[:80]
    filename = f"{safe_title or 'question_paper'}.pdf"
    return (pdf_bytes, filename)


def _build_sectioned_paper_from_text(combined_text: str, metadata: dict):
    """
    Create paper sections (MCQ/Short/Long/Case) strictly from combined_text.
    Returns sections list suitable for _build_question_paper_pdf.
    """
    import re

    def ensure_question_tag(q: str) -> str:
        q = (q or "").strip()
        q = re.sub(r"^\s*(\(?\d+[\).:-]\s+|\(?[a-zA-Z][\).:-]\s+|-+\s+)", "", q).strip()
        if not q:
            return q
        if re.match(r"^(what|how|explain|list|which|why|define|describe|discuss|compare)\b", q, re.I):
            return q
        low = q.lower()
        if "following" in low:
            return "List " + q
        if "how" in low:
            return "How " + q
        if any(k in low for k in ["explain", "describe", "discuss"]):
            return "Explain " + q
        if any(k in low for k in ["list", "name", "state"]):
            return "List " + q
        if "define" in low:
            return "What is " + q
        return "What " + q

    def sentence_pool(txt: str) -> list:
        sents = re.split(r"[.!?\n]+", txt or "")
        sents = [re.sub(r"\s+", " ", s).strip() for s in sents if 20 < len(s.strip()) < 360]
        seen = set()
        out = []
        for s in sents:
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(s)
        return out

    tpl = (metadata or {}).get("paperTemplate") or "auto"
    mode = (metadata or {}).get("mode") or "school"

    # Defaults for templates
    if tpl == "school_class8_science_50":
        mcq_n, short_n, long_n = 15, 10, 3  # 15*1 + 10*2 + 3*5 = 50
        mcq_m, short_m, long_m = 1, 2, 5
        title = f"Class 8 Science Question Paper"
        (metadata or {})["totalMarks"] = (metadata or {}).get("totalMarks") or 50
        (metadata or {})["examDuration"] = (metadata or {}).get("examDuration") or "2 Hours"
    elif tpl == "university_dbms_70":
        mcq_n, short_n, long_n = 10, 10, 5  # 10*1 + 10*2 + 5*10 = 70
        mcq_m, short_m, long_m = 1, 2, 10
        title = f"DBMS Question Paper (UG - 2nd Year)"
        (metadata or {})["totalMarks"] = (metadata or {}).get("totalMarks") or 70
        (metadata or {})["examDuration"] = (metadata or {}).get("examDuration") or "3 Hours"
    elif tpl == "university_os_70":
        mcq_n, short_n, long_n = 10, 10, 5  # same split, difficulty controlled via ordering below
        mcq_m, short_m, long_m = 1, 2, 10
        title = f"Operating System Question Paper (University)"
        (metadata or {})["totalMarks"] = (metadata or {}).get("totalMarks") or 70
        (metadata or {})["examDuration"] = (metadata or {}).get("examDuration") or "3 Hours"
    else:
        # auto: keep existing behavior by returning None
        return None, None

    pool = sentence_pool(combined_text)
    if not pool:
        return None, None

    # MCQs: use existing generator (already strict-from-text + tagged)
    mcqs = generate_questions_from_text(combined_text, num_questions=mcq_n)
    mcq_items = []
    for q in mcqs:
        mcq_items.append({
            "question": ensure_question_tag(q.get("question", "")),
            "options": q.get("options") or [],
            "correctAnswer": q.get("correctAnswer", 0),
            "marks": mcq_m
        })

    # Short: pull from sentences and ask tagged questions; answer is the sentence itself
    short_items = []
    cursor = 0
    while len(short_items) < short_n and cursor < len(pool):
        sent = pool[cursor]
        cursor += 1
        qtext = ensure_question_tag(sent)
        if not qtext.endswith("?"):
            qtext = qtext.rstrip(".:;") + "?"
        short_items.append({
            "question": qtext,
            "options": [],
            "answerText": sent,
            "marks": short_m,
            "marking": "Award marks for correct definition/explanation strictly from notes."
        })

    # Long: take longer sentences and ask explain/discuss; answer is 1-2 sentences from notes
    long_items = []
    while len(long_items) < long_n and cursor < len(pool):
        sent = pool[cursor]
        cursor += 1
        base = sent
        answer = base
        if cursor < len(pool) and len(answer) < 160:
            answer = answer + " " + pool[cursor]
            cursor += 1
        qtext = "Explain " + base
        qtext = ensure_question_tag(qtext)
        if not qtext.endswith("?"):
            qtext = qtext.rstrip(".:;") + "?"
        long_items.append({
            "question": qtext[:220],
            "options": [],
            "answerText": answer[:600],
            "marks": long_m,
            "marking": "Marks based on key points covered from the provided notes; no outside topics."
        })

    sections = [
        {
            "name": "Section A: MCQs",
            "instructions": "Attempt all questions. Choose the correct option.",
            "marksPerQuestion": mcq_m,
            "questions": mcq_items
        },
        {
            "name": "Section B: Short Answer Questions",
            "instructions": "Answer in brief. Strictly based on the provided notes.",
            "marksPerQuestion": short_m,
            "questions": short_items
        },
        {
            "name": "Section C: Long Answer Questions",
            "instructions": "Answer in detail. Strictly based on the provided notes.",
            "marksPerQuestion": long_m,
            "questions": long_items
        }
    ]

    return sections, title


@app.route('/generate-quiz-pdf', methods=['POST'])
def generate_quiz_pdf():
    # Accepts same form as /generate-quiz but returns a PDF of the quiz
    start_time = time.time()
    files = request.files.getlist('files')
    metadata_raw = request.form.get('metadata', '{}')
    try:
        metadata = json.loads(metadata_raw)
        num_questions = int(metadata.get('numberOfQuestions', 10))
    except:
        metadata = {}
        num_questions = 10

    try:
        texts = _read_file_texts(files, page_limit=10, force_ocr=bool(metadata.get('forceOcr', False)))
    except Exception as e:
        return jsonify({"error": "OCR failed", "detail": str(e)}), 500

    combined_text = " ".join(texts)[:10000]
    if not metadata.get('subject'):
        extracted_meta = extract_pdf_metadata(combined_text)
        if extracted_meta.get('subject'):
            metadata['subject'] = extracted_meta.get('subject')

    order_pref = (metadata.get('difficultyOrder') or 'low-to-high')
    if order_pref in ['low-to-high', 'high-to-low', 'mixed']:
        pool = generate_questions_from_text(combined_text, num_questions=max(30, num_questions*3))
        buckets = {'Easy': [], 'Medium': [], 'Hard': []}
        for q in pool:
            d = classify_difficulty(q.get('question') if isinstance(q, dict) else str(q))
            buckets[d].append(q)

        selected = []
        if order_pref == 'low-to-high':
            n_easy = max(1, int(num_questions * 0.4))
            n_medium = max(1, int(num_questions * 0.4))
            n_hard = num_questions - n_easy - n_medium
            selected.extend(buckets['Easy'][:n_easy])
            selected.extend(buckets['Medium'][:n_medium])
            selected.extend(buckets['Hard'][:n_hard])
        elif order_pref == 'high-to-low':
            n_hard = max(1, int(num_questions * 0.4))
            n_medium = max(1, int(num_questions * 0.4))
            n_easy = num_questions - n_hard - n_medium
            selected.extend(buckets['Hard'][:n_hard])
            selected.extend(buckets['Medium'][:n_medium])
            selected.extend(buckets['Easy'][:n_easy])
        else:
            seq = ['Easy', 'Medium', 'Hard']
            idx = 0
            while len(selected) < num_questions:
                level = seq[idx % 3]
                if buckets[level]:
                    selected.append(buckets[level].pop(0))
                else:
                    added = False
                    for l in ['Easy', 'Medium', 'Hard']:
                        if buckets[l]:
                            selected.append(buckets[l].pop(0))
                            added = True
                            break
                    if not added:
                        break
                idx += 1
        questions = selected[:num_questions]
    else:
        questions = generate_questions_from_text(combined_text, num_questions=num_questions)

    pdf_bytes, filename = _build_quiz_pdf(questions, title=(metadata.get('subject') or 'Quiz'), metadata=metadata)
    return (pdf_bytes, 200, { 'Content-Type': 'application/pdf', 'Content-Disposition': f'attachment; filename="{filename}"' })


@app.route('/generate-question-paper', methods=['POST'])
def generate_question_paper():
    # Creates a question paper from uploaded materials, arranging by increasing difficulty
    start_time = time.time()
    files = request.files.getlist('files')
    metadata_raw = request.form.get('metadata', '{}')
    try:
        metadata = json.loads(metadata_raw)
        total_questions = int(metadata.get('numberOfQuestions', 10))
    except:
        metadata = {}
        total_questions = 10

    try:
        texts = _read_file_texts(files, page_limit=20, force_ocr=bool(metadata.get('forceOcr', False)))
    except Exception as e:
        app.logger.error('OCR failed in generate_question_paper: %s', str(e))
        return jsonify({"error": "OCR failed", "detail": str(e)}), 500

    app.logger.info('generate-question-paper: extracted %d text blocks; combined length=%d', len(texts), len(" ".join(texts)))
    combined_text = " ".join(texts)[:20000]
    if not metadata.get('subject'):
        extracted_meta = extract_pdf_metadata(combined_text)
        if extracted_meta.get('subject'):
            metadata['subject'] = extracted_meta.get('subject')

    # If a template is requested, generate a proper sectioned paper (with answer key / marking scheme)
    def _to_bool(v, default=True):
        if v is None:
            return default
        if isinstance(v, bool):
            return v
        if isinstance(v, (int, float)):
            return bool(v)
        s = str(v).strip().lower()
        if s in ['true', '1', 'yes', 'y', 'on']:
            return True
        if s in ['false', '0', 'no', 'n', 'off']:
            return False
        return default

    include_answer_key = _to_bool(metadata.get('includeAnswerKey'), True)
    include_marking_scheme = _to_bool(metadata.get('includeMarkingScheme'), True)
    sections, paper_title = _build_sectioned_paper_from_text(combined_text, metadata)

    if sections and paper_title:
        pdf_bytes, filename = _build_question_paper_pdf(
            sections,
            title=(paper_title or (metadata.get('subject') or 'Question Paper')),
            metadata=metadata,
            include_answer_key=include_answer_key,
            include_marking_scheme=include_marking_scheme
        )
    else:
        # Backward-compatible: generate a larger pool of questions and order by difficulty
        pool = generate_questions_from_text(combined_text, num_questions=max(30, total_questions*3))
        app.logger.info('generate-question-paper: generated pool size=%d', len(pool))

        # If pool is empty, return an informative error so frontend does not receive a blank PDF
        if not pool or len(pool) == 0:
            app.logger.warning('No questions generated for question paper (empty pool)')
            return jsonify({"error": "No questions generated from the provided content. Try using clearer/scanned material or increase numberOfQuestions."}), 400

        # Classify into difficulty buckets
        buckets = {'Easy': [], 'Medium': [], 'Hard': []}
        for q in pool:
            d = classify_difficulty(q.get('question') if isinstance(q, dict) else str(q))
            buckets[d].append(q)

        ordered = []
        order_pref = (metadata.get('difficultyOrder') or 'low-to-high')

        if order_pref == 'low-to-high':
            # proportion: 40% easy, 40% medium, 20% hard
            n_easy = max(1, int(total_questions * 0.4))
            n_medium = max(1, int(total_questions * 0.4))
            n_hard = total_questions - n_easy - n_medium
            ordered.extend(buckets['Easy'][:n_easy])
            ordered.extend(buckets['Medium'][:n_medium])
            ordered.extend(buckets['Hard'][:n_hard])
        elif order_pref == 'high-to-low':
            # proportion: 40% hard, 40% medium, 20% easy
            n_hard = max(1, int(total_questions * 0.4))
            n_medium = max(1, int(total_questions * 0.4))
            n_easy = total_questions - n_hard - n_medium
            ordered.extend(buckets['Hard'][:n_hard])
            ordered.extend(buckets['Medium'][:n_medium])
            ordered.extend(buckets['Easy'][:n_easy])
        else:
            # mixed: round-robin from easy, medium, hard until filled
            seq = ['Easy', 'Medium', 'Hard']
            idx = 0
            while len(ordered) < total_questions:
                level = seq[idx % 3]
                if buckets[level]:
                    ordered.append(buckets[level].pop(0))
                else:
                    added = False
                    for l in ['Easy', 'Medium', 'Hard']:
                        if buckets[l]:
                            ordered.append(buckets[l].pop(0))
                            added = True
                            break
                    if not added:
                        break
                idx += 1

        ordered = ordered[:total_questions]
        app.logger.info('generate-question-paper: selected %d ordered questions', len(ordered))

        pdf_bytes, filename = _build_quiz_pdf(ordered, title=(metadata.get('subject') or 'Question Paper'), metadata=metadata)

    # Log pdf size and guard against empty/blank PDFs
    size = len(pdf_bytes) if pdf_bytes else 0
    app.logger.info('generate-question-paper: pdf size=%d bytes', size)
    if not pdf_bytes or size < 1000:
        app.logger.warning('Generated PDF is empty or too small - treating as no content')
        return jsonify({"error": "No questions were generated from the provided file. Try using clearer/scanned content or modify advanced options."}), 400

    return (pdf_bytes, 200, { 'Content-Type': 'application/pdf', 'Content-Disposition': f'attachment; filename="{filename}"' })


@app.route('/generate-quiz', methods=['POST'])
def generate_quiz():
    """Accept uploaded files and return generated questions as JSON."""
    start_time = time.time()
    files = request.files.getlist('files') or ([request.files.get('file')] if request.files.get('file') else [])
    metadata_raw = request.form.get('metadata', '{}')
    try:
        metadata = json.loads(metadata_raw)
    except Exception:
        metadata = {}

    try:
        texts = _read_file_texts(files, page_limit=20, force_ocr=bool(metadata.get('forceOcr', False)))
    except Exception as e:
        return jsonify({"error": "OCR failed", "detail": str(e)}), 500

    combined_text = " ".join(texts)[:20000]
    if not metadata.get('subject'):
        extracted_meta = extract_pdf_metadata(combined_text)
        if extracted_meta.get('subject'):
            metadata['subject'] = extracted_meta.get('subject')

    num_questions = int(metadata.get('numberOfQuestions', 10))

    # If difficulty ordering requested, generate a larger pool and select ordered questions
    order_pref = (metadata.get('difficultyOrder') or 'low-to-high')
    if order_pref in ['low-to-high', 'high-to-low', 'mixed']:
        pool = generate_questions_from_text(combined_text, num_questions=max(30, num_questions*3))
        buckets = {'Easy': [], 'Medium': [], 'Hard': []}
        for q in pool:
            d = classify_difficulty(q.get('question') if isinstance(q, dict) else str(q))
            buckets[d].append(q)

        selected = []
        if order_pref == 'low-to-high':
            n_easy = max(1, int(num_questions * 0.4))
            n_medium = max(1, int(num_questions * 0.4))
            n_hard = num_questions - n_easy - n_medium
            selected.extend(buckets['Easy'][:n_easy])
            selected.extend(buckets['Medium'][:n_medium])
            selected.extend(buckets['Hard'][:n_hard])
        elif order_pref == 'high-to-low':
            n_hard = max(1, int(num_questions * 0.4))
            n_medium = max(1, int(num_questions * 0.4))
            n_easy = num_questions - n_hard - n_medium
            selected.extend(buckets['Hard'][:n_hard])
            selected.extend(buckets['Medium'][:n_medium])
            selected.extend(buckets['Easy'][:n_easy])
        else:
            seq = ['Easy', 'Medium', 'Hard']
            idx = 0
            while len(selected) < num_questions:
                level = seq[idx % 3]
                if buckets[level]:
                    selected.append(buckets[level].pop(0))
                else:
                    added = False
                    for l in ['Easy', 'Medium', 'Hard']:
                        if buckets[l]:
                            selected.append(buckets[l].pop(0))
                            added = True
                            break
                    if not added:
                        break
                idx += 1
        questions = selected[:num_questions]
    else:
        questions = generate_questions_from_text(combined_text, num_questions=num_questions)

    return jsonify({
        "questions": questions,
        "processingTime": round(time.time() - start_time, 2),
        "metadata": metadata
    })

@app.route('/submit-quiz', methods=['POST'])
@login_required
def submit_quiz():
    """Track quiz submission and calculate score"""
    try:
        data = request.get_json()
        answers = data.get('answers', [])
        correct_count = data.get('correctCount', 0)
        total_questions = data.get('totalQuestions', 0)
        
        # Update user quiz stats
        current_user.quizzes_attempted = (current_user.quizzes_attempted or 0) + 1
        current_user.quizzes_completed = (current_user.quizzes_completed or 0) + 1
        db.session.commit()
        
        score_percentage = (correct_count / total_questions * 100) if total_questions > 0 else 0
        
        return jsonify({
            "success": True,
            "message": f"Quiz submitted successfully! You scored {correct_count}/{total_questions}",
            "score": score_percentage,
            "stats": {
                "quizzesGenerated": current_user.quizzes_generated or 0,
                "quizzesAttempted": current_user.quizzes_attempted or 0,
                "quizzesCompleted": current_user.quizzes_completed or 0
            }
        })
    except Exception as e:
        return jsonify({
            "error": "Failed to submit quiz",
            "detail": str(e)
        }), 500

# ==================== Database Initialization ====================

if __name__ == '__main__':
    # Initialize database before starting the app
    if models_available:
        try:
            with app.app_context():
                db.create_all()
                print("[OK] Database initialized successfully")
        except Exception as e:
            print(f"[WARN] Database initialization failed: {e}")
            print("[WARN] Continuing without database")
    else:
        print("[WARN] Running in analysis-only mode (no DB)")
    
    host = os.environ.get('HOST', '0.0.0.0')
    port = int(os.environ.get('PORT', 5000))
    print(f"[OK] Flask app starting on {host}:{port}")
    
    # Run the app
    app.run(host=host, port=port, debug=False, use_reloader=False)
